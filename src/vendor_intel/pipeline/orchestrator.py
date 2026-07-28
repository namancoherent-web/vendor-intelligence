"""Pipeline: Plan → Discovery → SSC enrich → Quality classify → CSV (~20 min)."""
from __future__ import annotations

import asyncio
import json
import os
import re
import time
from datetime import datetime, timezone
from typing import Any

from vendor_intel.config import Settings
from vendor_intel.enrichment.smart_enrichment import clear_enrichment_cache, enrich_companies
from vendor_intel.intelligence.classifier import classify_company
from vendor_intel.intelligence.signal_extractor import extract_signals
from vendor_intel.phase1.runner import run_phase1
from vendor_intel.phase2.discovery_fast import run_phase2_fast
from vendor_intel.pipeline.entity_gate import filter_companies, reject_reason
from vendor_intel.pipeline.llm_meter import get_meter, reset_meter
from vendor_intel.pipeline.plan_seeds import load_scope_from_plan, merge_seeds_first
from vendor_intel.pipeline.csv_fields import (
    extract_company_summary,
    extract_operational_presence,
    format_data_sources,
)
from vendor_intel.pipeline.quality_export import filter_for_export


def _build_query(query_context: dict[str, Any]) -> str:
    industry = str(query_context.get("industry") or "").strip()
    country = str(query_context.get("country") or "global").strip()
    functions = query_context.get("functions") or []
    fn = ", ".join(str(f) for f in functions[:5] if str(f).strip())
    parts = [p for p in [industry, f"in {country}" if country and country != "global" else "", fn] if p]
    return " ".join(parts).strip() or industry or "market companies"


def _industry_keywords(query_context: dict[str, Any], scope: dict[str, Any]) -> list[str]:
    from vendor_intel.pipeline.participant_domains import filter_industry_match_terms

    kws: list[str] = []
    kws.append(str(query_context.get("industry") or ""))
    kws.extend(str(f) for f in query_context.get("functions") or [])
    for field in ("relevance_keywords", "industry_terms"):
        kws.extend(str(x) for x in (scope.get(field) or []))
    out: list[str] = []
    seen: set[str] = set()
    for k in filter_industry_match_terms(
        [x.strip().lower() for x in kws if x and str(x).strip()]
    ):
        if k not in seen:
            seen.add(k)
            out.append(k)
    return out


def _profile_limits(
    settings: Settings, recall: bool, country: str | None = None
) -> tuple[int, int, float]:
    from vendor_intel.pipeline.geo_limits import pipeline_limits

    lim = pipeline_limits(settings, recall=recall, country=country)
    return int(lim["discover"]), int(lim["enrich"]), float(lim["min_conf"])


def _enrichment_source(smart_data: dict[str, Any]) -> str:
    if not smart_data:
        return "none"
    if smart_data.get("source"):
        return str(smart_data["source"])
    if smart_data.get("error"):
        return f"error:{smart_data.get('error')}"[:40]
    return "smart_crawl"


def _website_display(domain: str, smart_data: dict[str, Any]) -> str:
    dom = (domain or "").strip()
    if not dom:
        return ""
    data = smart_data.get("data") or {}
    if isinstance(data, dict):
        co = data.get("company") or {}
        if isinstance(co, dict) and co.get("website"):
            return str(co["website"])
    return dom if dom.startswith("http") else f"https://{dom}"


_REJECT_FRIENDLY: dict[str, str] = {
    "not_relevant": "the site did not clearly mention this market - relevance could not be confirmed",
    "non_product_site": "not a company/product site (directory, blog, marketplace, or aggregator)",
    "market_research_site": "a market-research / news / data site, not a market participant",
    "consulting_not_participant": "appears to be consulting/advisory, not a market participant",
    "below_quality_threshold": "too little evidence on the site to confirm",
    "weak_product_fit": "products shown on the site did not match the market closely enough",
    "off_topic_pharma": "site content pointed to a different (pharma) domain",
    "off_topic_pharma_cdmo": "site content pointed to a different (pharma) domain",
    "duplicate_domain": "duplicate of another listed company",
}


def _friendly_reject(code: str) -> str:
    code = (code or "").strip()
    if not code:
        return "could not confirm this company operates in the market"
    if code.startswith("geo_mismatch"):
        return (
            f"headquartered in {code.split(':', 1)[-1]} with no detected operations in the "
            f"target geography"
        )
    return _REJECT_FRIENDLY.get(code, f"could not confirm ({code})")


def _verification_status(v: dict[str, Any]) -> tuple[bool, str]:
    """Is an exported row actually substantiated? Reason if not.

    Seeds / known players (we already resolved their official domain) are NEVER demoted to
    'Not Verified' for a mere crawl failure — a transient fetch error must not banish a
    recognised major. They only leave the list if classified clearly off-market.
    """
    src = str(v.get("enrichment_source") or "").lower()
    crawl_failed = (not src) or src in ("none", "no_crawl") or src.startswith("error") or "fail" in src
    if not v.get("is_relevant"):
        return False, _friendly_reject(str(v.get("reject_reason") or ""))
    if v.get("is_seed"):
        return True, ""  # known/pinned player with a resolved domain — keep it
    if crawl_failed:
        return (
            False,
            "website could not be retrieved - listed by name only; its site did not confirm "
            "activity in this market",
        )
    return True, ""


async def run_pipeline(
    query_context: dict[str, Any],
    settings: Settings | None = None,
    *,
    enrich_limit: int | None = None,
    classify_limit: int | None = None,
) -> dict[str, Any]:
    settings = settings or Settings.load()
    recall = bool(getattr(settings, "pipeline_recall_mode", False))
    profile = str(getattr(settings, "pipeline_profile", "quality") or "quality")
    quality = not recall and profile in ("quality", "balanced")
    country = str(query_context.get("country") or "global")
    from vendor_intel.pipeline.geo_limits import is_global_geography, pipeline_limits

    lim = pipeline_limits(settings, recall=recall, country=country)
    discover_cap, enrich_cap, min_conf = (
        int(lim["discover"]),
        int(lim["enrich"]),
        float(lim["min_conf"]),
    )
    classify_cap = classify_limit if classify_limit is not None else discover_cap
    enrich_cap = enrich_limit if enrich_limit is not None else enrich_cap
    export_max = int(lim["export_max"])
    export_min = int(lim["export_min"])
    min_quality = float(lim["min_quality"])
    geo_label = "global (higher volume)" if is_global_geography(country) else country

    reset_meter()
    t0 = time.perf_counter()
    query = _build_query(query_context)
    print(f"\n=== Pipeline start ===\n  Query: {query}", flush=True)
    print(
        f"  Profile: {profile}" + (" (recall/noisy)" if recall else " (quality)"),
        flush=True,
    )
    print(
        f"  Geography: {geo_label} - discover<={discover_cap}, enrich<={enrich_cap}, "
        f"export up to {export_max} rows (no padding)",
        flush=True,
    )
    use_ssc = bool(getattr(settings, "pipeline_use_ssc", True)) and profile != "deep"
    print(f"  Enrichment: {'SSC server-side' if use_ssc else 'smart_crawl'}", flush=True)

    clear_enrichment_cache()
    from vendor_intel.pipeline.cancel import check_cancelled

    # When the user specifies functional categories, steer discovery toward them:
    # the market-understanding LLM generates prompts for these company types only.
    focus_sections = [
        str(s).strip()
        for s in (query_context.get("sections") or [])
        if str(s).strip() and str(s).strip().lower() != "other"
    ]
    plan_query = query
    if focus_sections:
        plan_query = (
            f"{query} -- only include these company functions / types: "
            f"{'; '.join(focus_sections)}"
        )
        print(
            f"  [pipeline] discovery focused on {len(focus_sections)} requested company types",
            flush=True,
        )

    print("  [pipeline] Phase 1 — query plan", flush=True)
    check_cancelled("before Phase 1")
    plan = await run_phase1(plan_query, settings)
    check_cancelled("after Phase 1")
    get_meter().add_phase1()
    plan_path = plan.get("plan_path") or plan.get("output_path")
    scope = load_scope_from_plan(str(plan_path) if plan_path else None)
    industry_kws = _industry_keywords(query_context, scope)

    # Diagnostic: SEED AUDIT mode (SEED_AUDIT=1 or --seed-audit). Seeds are NOT pinned into the
    # run — the system discovers purely on its own — and at the end we report which seed companies
    # it found / kept / rejected. Temporary, for understanding the code's true recall.
    seed_audit = os.getenv("SEED_AUDIT", "").strip().lower() in ("1", "true", "yes", "on")
    audit_seed_names = [
        str(s).strip() for s in (query_context.get("seed_companies") or []) if str(s).strip()
    ]

    # Analyst-provided seed companies: resolve domains and guarantee them into the run.
    user_seed_names = audit_seed_names
    unresolved_seed_names: list[str] = []
    seed_domain_map: dict[str, str] = {}  # normalized seed name -> resolved domain
    if user_seed_names and not seed_audit:
        from vendor_intel.pipeline.plan_seeds import (
            merge_user_seeds_into_scope,
            resolve_user_seeds,
        )

        resolved_seeds, unresolved_seed_names = await resolve_user_seeds(
            user_seed_names, settings, country
        )
        seed_domain_map = {
            re.sub(r"[^a-z0-9]", "", str(s.get("canonical_name") or "").lower()): str(s.get("primary_domain") or "")
            for s in resolved_seeds
        }
        scope = merge_user_seeds_into_scope(scope, resolved_seeds)
        print(
            f"  [pipeline] {len(resolved_seeds)}/{len(user_seed_names)} analyst seeds "
            f"resolved and pinned into the run",
            flush=True,
        )
    elif seed_audit:
        print(
            f"  [pipeline] SEED AUDIT mode — {len(audit_seed_names)} curated seeds NOT pinned "
            f"(LLM enumeration still active); measuring whether the system surfaces them on its own",
            flush=True,
        )

    # Discovered candidates from ALL non-curated sources (LLM enumeration, directory mining,
    # Wikidata). These are LEADS, not trusted seeds: each is added with is_seed=False so it must
    # pass crawl + classification on its own evidence. Only the analyst's CURATED list is pinned.
    discovered_candidates: list[dict] = []
    authoritative_domains: set[str] = set()

    # Widen the REAL-company base from the LLM's own knowledge (what GPT/Google would list) —
    # the brand-majors plain search misses. Runs even in seed-audit mode (the audit only unpins
    # the curated seed FILE), so the audit measures whether the system surfaces majors on its own.
    if not recall and getattr(settings, "pipeline_enumerate_players", True):
        try:
            from vendor_intel.clients.claude import ClaudeClient
            from vendor_intel.funnel.market_understanding import enumerate_market_players
            from vendor_intel.pipeline.plan_seeds import merge_user_seeds_into_scope

            section_hints = [
                str(s).strip() for s in (query_context.get("sections") or []) if str(s).strip()
            ]
            players = enumerate_market_players(
                str(scope.get("market") or query_context.get("industry") or ""),
                country,
                section_hints,
                ClaudeClient(settings),
                settings,
                market_definition=str(scope.get("market_definition") or ""),
            )
            if players:
                # Cap how many LLM leads enter the run. Env-configurable (PIPELINE_MAX_LLM_PLAYERS)
                # so a run leans on WEB/structured discovery and keeps LLM a SMALL base (e.g. 30).
                MAX_LLM_PLAYERS = int(os.getenv("PIPELINE_MAX_LLM_PLAYERS", "250") or "250")
                if len(players) > MAX_LLM_PLAYERS:
                    _maj = [p for p in players if p.get("is_major")]
                    _oth = [p for p in players if not p.get("is_major")]
                    players = (_maj + _oth)[:MAX_LLM_PLAYERS]
                # AI-enumerated names are LEADS, not trusted seeds. Add as candidates (is_seed=False)
                # so each must pass crawl + classification on its own evidence — the AI can SUGGEST a
                # company but can no longer FORCE-keep it. (No major_player_names: we deliberately do
                # NOT let the LLM's "is_major" claim auto-promote a row past the relevance check.)
                for p in players:
                    dom = str(p.get("primary_domain") or "").lower().removeprefix("www.")
                    if dom and "." in dom:
                        discovered_candidates.append({
                            "name": str(p.get("canonical_name") or ""),
                            "domain": dom,
                            "company_function": str(p.get("company_function") or ""),
                            "discovery_source": "llm_enumeration",
                            "is_seed": False,
                        })
                get_meter().add_market_understanding()
                print(
                    f"  [pipeline] enumerated {len(players)} LLM leads (cap {MAX_LLM_PLAYERS}) — added "
                    f"as candidates that must pass classification (not auto-kept)",
                    flush=True,
                )
        except Exception as exc:
            print(f"  [pipeline] player enumeration skipped: {exc}", flush=True)

    # Directory / listicle / association / exhibitor mining: harvest real companies NAMED on
    # list pages (the source pages themselves are never added — only the companies on them).
    if not recall and getattr(settings, "pipeline_directory_mining", True):
        try:
            from vendor_intel.clients.claude import ClaudeClient
            from vendor_intel.discovery.directory_mining import mine_directories

            mkt = str(scope.get("market") or query_context.get("industry") or "")
            sec_hints = [str(s).strip() for s in (query_context.get("sections") or []) if str(s).strip()]
            ind_terms = [str(t).strip() for t in (scope.get("industry_terms") or []) if str(t).strip()]
            mined = mine_directories(
                mkt, country, sec_hints, settings, ClaudeClient(settings), industry_terms=ind_terms
            )
            if mined:
                from vendor_intel.discovery.entity_extract import is_blocked_domain, is_listicle_domain
                from vendor_intel.pipeline.plan_seeds import resolve_user_seeds

                def _cd(d: str) -> str:
                    d = str(d or "").strip().lower()
                    return d.removeprefix("http://").removeprefix("https://").removeprefix("www.").split("/")[0]

                # Companies the directory linked directly → trust that domain (no guessing).
                linked: list[dict] = []
                to_resolve: list[str] = []
                for m in mined:
                    dom = _cd(m.get("domain"))
                    if dom and "." in dom and not is_blocked_domain(dom) and not is_listicle_domain(dom):
                        linked.append({"canonical_name": m.get("name"), "primary_domain": dom, "company_function": ""})
                    else:
                        to_resolve.append(str(m.get("name") or ""))
                # Names the directory didn't link → resolve robustly (LLM + DNS + search fallback).
                # CAP the count: a low-quality mega-directory (e.g. an ISP-license list) can dump
                # hundreds of tiny names, and resolving each via web search floods the search
                # backends (the timeouts you saw). Directory-LINKED companies are unaffected.
                # PRIORITISE the most company-like names first, THEN cap — a real major (which the
                # LLM resolves instantly) must never be dropped just because the mega-list happened
                # to place it after position 60. The obscure tail (which falls to slow web-search
                # resolution) is what gets deferred instead. This was the main cause of top vendors
                # like Abbott / Bio-Rad going missing on a broad market.
                resolved: list[dict] = []
                if to_resolve:
                    _corp_suffix = re.compile(
                        r"\b(inc|ltd|llc|corp|corporation|gmbh|ag|sa|plc|co|company|laboratories|"
                        r"labs|diagnostics|biosystems|technologies|scientific|genomics|bioscience|"
                        r"biosciences|sciences|healthcare|medical|systems|group|therapeutics|nv|"
                        r"pte|sas|kgaa|holding|holdings)\b", re.I,
                    )

                    def _resolve_rank(nm: str) -> tuple:
                        s = (nm or "").strip()
                        words = s.split()
                        capd = sum(1 for w in words if w[:1].isupper())
                        score = (
                            (2 if _corp_suffix.search(s) else 0)
                            + (1 if 2 <= len(s) <= 45 else 0)
                            + (1 if 1 <= len(words) <= 5 else 0)
                            + (1 if words and capd / len(words) >= 0.5 else 0)
                        )
                        return (-score, len(s))  # highest score first, then shorter names

                    to_resolve.sort(key=_resolve_rank)
                    _CAP = 90
                    if len(to_resolve) > _CAP:
                        print(f"  [directory] resolving the top {_CAP} most company-like names "
                              f"(of {len(to_resolve)} unlinked); obscure tail deferred", flush=True)
                    resolved, _unres = await resolve_user_seeds(to_resolve[:_CAP], settings, country)
                    resolved = [
                        r for r in resolved
                        if _cd(r.get("primary_domain"))
                        and not is_blocked_domain(_cd(r.get("primary_domain")))
                        and not is_listicle_domain(_cd(r.get("primary_domain")))
                    ]
                all_dir = linked + resolved
                for r in all_dir:
                    dom = _cd(r.get("primary_domain"))
                    if dom and "." in dom:
                        discovered_candidates.append({
                            "name": str(r.get("canonical_name") or ""),
                            "domain": dom,
                            "company_function": "",
                            "discovery_source": "directory",
                            "is_seed": False,
                        })
                print(
                    f"  [pipeline] directory mining added {len(all_dir)} candidate(s) from list pages "
                    f"({len(linked)} via the directory's own link, no guess) — must pass classification",
                    flush=True,
                )
        except Exception as exc:
            print(f"  [pipeline] directory mining skipped: {exc}", flush=True)

    # Wikidata structured discovery: real organizations from Wikidata (official websites),
    # NOT LLM recall. Surfaces the regional/obscure tail that never ranks in search or sits on a
    # listicle. Added as DISCOVERED CANDIDATES (is_seed=False) so they still face relevance
    # classification — that drops gov/agency noise (NASA, DARPA) semantically. Their domains ARE
    # authoritative, so they bypass the entity-gate's blunt rules (e.g. a real .gov-hosted
    # operator like NIGCOMSAT isn't killed by the gov-TLD block).
    if not recall and getattr(settings, "pipeline_wikidata", True):
        try:
            from vendor_intel.discovery.entity_extract import is_blocked_domain, is_listicle_domain
            from vendor_intel.discovery.wikidata_discovery import discover_via_wikidata

            def _cdw(d: str) -> str:
                d = str(d or "").strip().lower()
                return d.removeprefix("http://").removeprefix("https://").removeprefix("www.").split("/")[0]

            mkt = str(scope.get("market") or query_context.get("industry") or "")
            sec_hints = [str(s).strip() for s in (query_context.get("sections") or []) if str(s).strip()]
            ind_terms = [str(t).strip() for t in (scope.get("industry_terms") or []) if str(t).strip()]
            wd_rows = await asyncio.to_thread(discover_via_wikidata, mkt, sec_hints, ind_terms)
            wd_n = 0
            for r in wd_rows:
                dom = _cdw(r.get("domain"))
                if not dom or "." not in dom or is_blocked_domain(dom) or is_listicle_domain(dom):
                    continue
                discovered_candidates.append(
                    {"name": str(r.get("name") or ""), "domain": dom, "company_function": "",
                     "discovery_source": "wikidata", "is_seed": False}
                )
                authoritative_domains.add(dom)  # structured-source domains bypass the gate's fuzzy rules
                wd_n += 1
            if wd_n:
                print(f"  [pipeline] Wikidata added {wd_n} candidate(s) from structured data (official websites, no guess)", flush=True)
        except Exception as exc:
            print(f"  [pipeline] Wikidata discovery skipped: {exc}", flush=True)

    # Persistent cross-run discovery store: re-inject everything this market has discovered before
    # as candidates (is_seed=False) so they're re-crawled + re-classified. Turns per-run volatility
    # into cumulative coverage — each run unions its finds into the store afterwards.
    store_market = str(scope.get("market") or query_context.get("industry") or "")
    store_by_dom: dict[str, dict] = {}  # normalized domain -> stored record (rich fields)
    if not recall:
        try:
            from vendor_intel.pipeline.discovery_store import enabled as _store_on
            from vendor_intel.pipeline.discovery_store import load_discovered

            if _store_on():
                stored = load_discovered(store_market, country)
                print(f"  [pipeline] discovery store: market={store_market!r} geo={country!r} -> "
                      f"loaded {len(stored)} previously-found compan(ies)", flush=True)
                for r in stored:
                    d = str(r.get("domain") or "").lower().removeprefix("www.")
                    if not d:
                        continue
                    store_by_dom[d] = r
                    discovered_candidates.append({
                        "name": r["name"], "domain": r["domain"], "company_function": "",
                        "discovery_source": "store", "is_seed": False,
                    })
                    # Previously-CONFIRMED companies for THIS market+geo are trusted by provenance:
                    # bypass the gate's fuzzy rejects (non_product_site / gov_or_edu) that misfire on
                    # real operators. The store is now scoped per geography, so this can no longer
                    # leak a different region's confirmed companies into this run.
                    authoritative_domains.add(r["domain"])
        except Exception as exc:
            import traceback
            print(f"  [pipeline] discovery store load FAILED: {exc}", flush=True)
            traceback.print_exc()

    # Partner / reseller page mining: the regional integrator tail (Q-KON, iSAT Africa, Paratus,
    # Satcom Networks Africa, X2nSat) is listed as PARTNERS of the majors. Anchor on known players
    # and read their partner pages — reads real partner directories, not LLM recall.
    if not recall and getattr(settings, "pipeline_partner_mining", True):
        try:
            from vendor_intel.clients.claude import ClaudeClient
            from vendor_intel.discovery.directory_mining import mine_partner_pages
            from vendor_intel.discovery.entity_extract import is_blocked_domain, is_listicle_domain
            from vendor_intel.pipeline.plan_seeds import resolve_user_seeds

            def _cdp(d: str) -> str:
                d = str(d or "").strip().lower()
                return d.removeprefix("http://").removeprefix("https://").removeprefix("www.").split("/")[0]

            anchor_names = [str(s.get("canonical_name") or "").strip()
                            for s in (scope.get("seed_companies") or []) if isinstance(s, dict)]
            anchor_names += [r.get("name", "") for r in list(store_by_dom.values())[:12]]
            anchor_names = [a for a in dict.fromkeys(anchor_names) if a][:10]
            mkt = str(scope.get("market") or query_context.get("industry") or "")
            prows = await asyncio.to_thread(
                mine_partner_pages, anchor_names, mkt, country, settings, ClaudeClient(settings)
            )
            linked = [r for r in prows if _cdp(r.get("domain")) and "." in _cdp(r.get("domain"))]
            unlinked = [str(r.get("name") or "") for r in prows
                        if not (_cdp(r.get("domain")) and "." in _cdp(r.get("domain")))]
            resolved: list[dict] = []
            if unlinked:
                resolved, _u = await resolve_user_seeds(unlinked[:40], settings, country)
            pn = 0
            for nm, dom in (
                [(r.get("name"), _cdp(r.get("domain"))) for r in linked]
                + [(r.get("canonical_name"), _cdp(r.get("primary_domain"))) for r in resolved]
            ):
                if dom and "." in dom and not is_blocked_domain(dom) and not is_listicle_domain(dom):
                    discovered_candidates.append({
                        "name": str(nm or ""), "domain": dom, "company_function": "",
                        "discovery_source": "partner", "is_seed": False,
                    })
                    pn += 1
            if pn:
                print(f"  [pipeline] partner mining added {pn} candidate(s) from partner/reseller pages", flush=True)
        except Exception as exc:
            print(f"  [pipeline] partner mining skipped: {exc}", flush=True)

    print("  [pipeline] Phase 2 — discovery + entity gate", flush=True)
    check_cancelled("before Phase 2")
    p2 = await run_phase2_fast(query, settings, phase1_plan_path=plan_path)
    check_cancelled("after Phase 2")
    companies: list[dict[str, str]] = list(p2.get("companies") or [])
    companies, seed_doms = merge_seeds_first(companies, scope)
    # Inject ALL discovered candidates (LLM leads + directory + Wikidata), deduped by domain —
    # discovered, NOT pinned: every one must earn its place through crawl + classification.
    if discovered_candidates:
        have_dom = {str(c.get("domain") or "").lower().removeprefix("www.") for c in companies}
        add_cand = []
        for w in discovered_candidates:
            d = str(w.get("domain") or "").lower().removeprefix("www.")
            if d and d not in have_dom:
                have_dom.add(d)
                add_cand.append(w)
        companies += add_cand
        if add_cand:
            print(f"  [pipeline] injected {len(add_cand)} discovered candidate(s) "
                  f"(LLM/directory/Wikidata) as evidence-judged, not pinned", flush=True)
    print(f"  [pipeline] {len(companies)} after seeds merge ({len(seed_doms)} seeds)", flush=True)

    # Independent-recall signal for the run summary. ALL curated seeds stay pinned in the final;
    # here we only MEASURE how many of them the system would have surfaced on its own — a pinned
    # seed whose domain was ALSO produced by a non-seed channel (web search, LLM enumeration,
    # directory, Wikidata, discovery store, partner mining).
    def _nd_recall(d: str) -> str:
        d = str(d or "").strip().lower()
        return d.removeprefix("http://").removeprefix("https://").removeprefix("www.").split("/")[0]

    _seed_dom_to_name: dict[str, str] = {}
    for _c in companies:
        if _c.get("is_seed"):
            _sd = _nd_recall(_c.get("domain"))
            if _sd:
                _seed_dom_to_name.setdefault(_sd, str(_c.get("name") or _c.get("company") or _sd))
    _nonseed_doms = {_nd_recall(w.get("domain")) for w in discovered_candidates}
    _nonseed_doms |= {_nd_recall(c.get("domain")) for c in companies if not c.get("is_seed")}
    _nonseed_doms.discard("")
    seeds_found_independently = sorted(
        name for dom, name in _seed_dom_to_name.items() if dom in _nonseed_doms
    )
    seeds_pinned_total = len(_seed_dom_to_name)
    if seeds_pinned_total:
        print(
            f"  [pipeline] independent recall: system surfaced "
            f"{len(seeds_found_independently)}/{seeds_pinned_total} curated seeds on its own",
            flush=True,
        )

    if not recall:
        # Trusted past the gate's fuzzy rules = analyst's CURATED seeds + AUTHORITATIVE-source
        # (Wikidata) companies — NOT LLM-enumerated guesses (so NASA/GPS still face the gate).
        curated_domains = {d for d in seed_domain_map.values() if d} | authoritative_domains
        companies, rejected_p2 = filter_companies(companies, trusted_domains=curated_domains)
        if rejected_p2:
            print(f"  [pipeline] entity_gate removed {len(rejected_p2)} non-participants", flush=True)

    # COST SAVER: store companies were already CONFIRMED relevant in a prior run. By default, carry
    # them straight to the output instead of re-crawling + re-classifying all of them every run —
    # so only FRESH discoveries incur crawl + LLM cost. Set PIPELINE_STORE_REVALIDATE=1 to re-check
    # them (more thorough, more expensive). Store companies are store-only by construction (a fresh
    # source would have claimed the domain before the store was injected).
    store_carry: list[dict[str, str]] = []
    if (
        not recall
        and store_by_dom
        # OPT-IN cost-saver: default OFF so store companies are re-crawled + re-classified normally
        # (correctness first). Set PIPELINE_STORE_CHEAP=1 to carry them without re-processing.
        and os.getenv("PIPELINE_STORE_CHEAP", "0").strip().lower() in ("1", "true", "yes", "on")
    ):
        def _nd(c):
            return str(c.get("domain") or "").lower().removeprefix("www.")
        # Carry ANY already-confirmed company (whatever source re-surfaced it), except pinned seeds.
        store_carry = [c for c in companies if not c.get("is_seed") and _nd(c) in store_by_dom]
        if store_carry:
            companies = [c for c in companies if c.get("is_seed") or _nd(c) not in store_by_dom]
            print(f"  [pipeline] cost-saver: carrying {len(store_carry)} already-confirmed compan(ies) "
                  f"to output WITHOUT re-crawl/classify (set PIPELINE_STORE_REVALIDATE=1 to re-check)", flush=True)

    cleaned: list[dict[str, str]] = []
    seen_dom: set[str] = set()
    from vendor_intel.pipeline.quality_export import _export_domain

    for c in companies:
        dom = _export_domain(c)
        name = (c.get("name") or "").strip()
        if not dom or not name or dom in seen_dom:
            continue
        seen_dom.add(dom)
        cleaned.append(c)

    from vendor_intel.utils.domain_corrections import fix_company_list

    top = fix_company_list(cleaned[:classify_cap], country=country, recall_mode=recall)
    # DNS/alias fix can change host (e.g. braskem.com → www.braskem.com.br) — refresh seed_doms for export
    seed_doms = {
        (c.get("domain") or "").strip().lower()
        for c in top
        if c.get("is_seed") and (c.get("domain") or "").strip()
    }
    # Off-market pre-filter: drop candidates whose name/domain clearly signals a
    # different industry (scope-driven exclude/negative keywords) BEFORE enrichment,
    # so we don't waste crawls on obvious junk. Seeds are exempt.
    if not recall:
        from vendor_intel.pipeline.market_relevance import exclude_keyword_hit, keyword_profile

        prof = keyword_profile(scope, query_context)
        kept_top: list[dict[str, str]] = []
        off_market = 0
        for c in top:
            nm = (c.get("name") or "").strip()
            dom = (c.get("domain") or "").strip()
            if c.get("is_seed") or dom in seed_doms:
                kept_top.append(c)
                continue
            if exclude_keyword_hit(f"{nm} {dom}", dom, prof, name=nm):
                off_market += 1
                continue
            kept_top.append(c)
        if off_market:
            top = kept_top
            print(
                f"  [pipeline] off-market pre-filter dropped {off_market} before enrichment",
                flush=True,
            )
    print(f"  [pipeline] {len(top)} companies to process", flush=True)

    print("  [pipeline] Phase 3 — enrichment", flush=True)
    check_cancelled("before enrichment")
    effective_enrich = min(len(top), max(enrich_cap, discover_cap))
    enrich_batch = top[:effective_enrich]
    enrich_concurrent = int(getattr(settings, "pipeline_enrich_concurrent", 8) or 8)
    enriched = await enrich_companies(
        enrich_batch,
        limit=effective_enrich,
        max_concurrent=enrich_concurrent,
        country=country,
        use_ssc=use_ssc,
    )
    check_cancelled("after enrichment")

    print("  [pipeline] Phase 4 — classification (quality)", flush=True)
    check_cancelled("before classification")
    from vendor_intel.clients.claude import ClaudeClient

    client = ClaudeClient(settings)
    from vendor_intel.pipeline.sections import build_section_taxonomy, main_product_label

    main_product = main_product_label(query_context, scope)
    custom_sections = [
        str(s).strip() for s in (query_context.get("sections") or []) if str(s).strip()
    ]
    if custom_sections:
        if not any(s.lower() == "other" for s in custom_sections):
            custom_sections = custom_sections + ["Other"]
        section_taxonomy = custom_sections
        print(f"  [pipeline] using {len(custom_sections)} custom CEO sections", flush=True)
    else:
        section_taxonomy = build_section_taxonomy(main_product)
    classify_ctx = {
        **query_context,
        "plan_keywords": industry_kws,
        "scope": scope,
        "market": str(scope.get("market") or query_context.get("industry") or ""),
        "main_product": main_product,
        "value_chain_sections": section_taxonomy,
    }
    classify_concurrent = int(getattr(settings, "pipeline_classify_concurrent", 8) or 8)
    print(
        f"  [pipeline] classifying {len(top)} companies "
        f"(concurrent={classify_concurrent})",
        flush=True,
    )
    sem = asyncio.Semaphore(max(1, classify_concurrent))
    done = 0

    async def _classify_one(c: dict[str, str]) -> dict[str, Any]:
        nonlocal done
        name = c["name"]
        dom = c["domain"]
        smart_data = enriched.get(name) or enriched.get(dom) or {"error": "no_crawl", "data": {}}
        is_seed = dom in seed_doms or bool(c.get("is_seed"))
        async with sem:
            check_cancelled("during classification")
            signals = extract_signals(smart_data, industry_keywords=industry_kws)
            verdict = await classify_company(
                c,
                smart_data,
                classify_ctx,
                settings=settings,
                client=client,
                recall_mode=recall,
                quality_mode=quality,
                is_seed=is_seed,
            )
        supplemented = verdict.pop("_enriched", None)
        classify_signals = verdict.pop("_classify_signals", None)
        if supplemented is not None:
            smart_data = supplemented
            enriched[name] = supplemented
            enriched[dom] = supplemented
        done += 1
        if done % 10 == 0 or done == len(top):
            print(f"  [pipeline] classified {done}/{len(top)}", flush=True)
        verdict["signals"] = classify_signals or extract_signals(
            smart_data, industry_keywords=industry_kws
        )
        verdict["discovery_source"] = c.get("discovery_source", "search")
        verdict["enrichment_source"] = _enrichment_source(smart_data)
        verdict["website"] = _website_display(dom, smart_data)
        verdict["is_seed"] = is_seed
        verdict["company_function"] = (
            c.get("company_function") or verdict.get("company_function") or ""
        )
        # Summary comes from the LLM's clean prose field; the CSV layer composes a
        # clean fallback from role/products. We deliberately do NOT use raw scraped
        # page text here (it leaks nav menus, cookie banners, markdown, mojibake).
        verdict["company_summary"] = str(verdict.get("summary") or "")
        verdict["operational_presence"] = extract_operational_presence(
            smart_data,
            str(query_context.get("country") or ""),
            signals=signals,
        )
        verdict["data_sources"] = format_data_sources(
            str(c.get("discovery_source") or verdict.get("discovery_source") or ""),
            verdict["enrichment_source"],
            smart_data,
        )
        return verdict

    classified = list(await asyncio.gather(*[_classify_one(c) for c in top]))

    # Honor analyst-assigned sections for must-haves ("Name | Section" in the seed file):
    # the pinned company is placed in that exact section regardless of keyword routing.
    seed_sections = query_context.get("seed_sections") or {}
    if seed_sections:
        import re as _re

        def _ck(s: str) -> str:
            return _re.sub(r"[^a-z0-9]", "", str(s or "").lower())

        smap = {_ck(k): v for k, v in seed_sections.items() if v}
        for v in classified:
            cn = _ck(v.get("company") or v.get("brand") or "")
            if not cn:
                continue
            for sk, sec in smap.items():
                if sk and (sk in cn or cn in sk):
                    v["_forced_section"] = sec
                    break

    if recall:
        final = [v for v in classified if float(v.get("confidence") or 0) >= min_conf]
        if not final:
            final = list(classified)
        export_rejected: list[dict[str, Any]] = []
    else:
        from collections import Counter

        final, export_rejected = filter_for_export(
            classified,
            enriched,
            query_context,
            scope=scope,
            seed_domains=seed_doms,
            min_rows=export_min,
            max_rows=export_max,
            min_quality=min_quality,
            pad_to_min=False,
        )
        if export_rejected:
            reasons = Counter(
                str(r.get("export_reject") or "?") for r in export_rejected
            )
            top = ", ".join(f"{k}={v}" for k, v in reasons.most_common(4))
            print(f"  [pipeline] export rejects ({len(export_rejected)}): {top}", flush=True)
        roles = Counter(str(r.get("role") or "?") for r in final) if final else Counter()
        if roles:
            print(
                f"  [pipeline] export roles: "
                + ", ".join(f"{k}={v}" for k, v in roles.most_common()),
                flush=True,
            )

    if getattr(settings, "pipeline_shuffle_export", False):
        # Stable pseudo-random order (no quality clustering / seeds-first tell).
        import hashlib

        final.sort(
            key=lambda x: hashlib.md5(
                str(x.get("domain") or x.get("website") or x.get("company") or "").encode()
            ).hexdigest()
        )
    else:
        final.sort(
            key=lambda x: float(x.get("quality_score") or x.get("confidence") or 0), reverse=True
        )

    from vendor_intel.pipeline.quality_export import dedupe_export_rows

    before_dedupe = len(final)
    final = dedupe_export_rows(final)
    if before_dedupe != len(final):
        print(
            f"  [pipeline] deduped export rows: {before_dedupe} -> {len(final)}",
            flush=True,
        )

    # Strict geography: drop companies clearly headquartered outside the target region
    # (seeds exempt; companies with no geo signal are kept). Disabled via PIPELINE_STRICT_GEO
    # / --no-geo-filter to keep global players that operate in the target geography.
    if not recall and getattr(settings, "pipeline_strict_geo", True):
        from vendor_intel.pipeline.geo_filter import geo_mismatch_reason, resolve_target_countries

        if resolve_target_countries(country):
            kept: list[dict[str, Any]] = []
            geo_dropped: list[dict[str, Any]] = []
            for r in final:
                reason = geo_mismatch_reason(r, r.get("signals"), country)
                if reason:
                    r["export_reject"] = reason
                    geo_dropped.append(r)
                else:
                    kept.append(r)
            if geo_dropped:
                final = kept
                export_rejected = list(export_rejected) + geo_dropped
                from collections import Counter

                by_country = Counter(
                    str(r.get("export_reject") or "").split(":", 1)[-1] for r in geo_dropped
                )
                print(
                    f"  [pipeline] strict geo ({country}) dropped {len(geo_dropped)}: "
                    + ", ".join(f"{k}={v}" for k, v in by_country.most_common(5)),
                    flush=True,
                )

    # Confirmation pass: only keep companies the system actually substantiated in the
    # main list. Named/known companies it could NOT confirm are reported separately with
    # a reason (no website found, site unreachable, or relevance not confirmed).
    unverified_companies: list[dict[str, Any]] = []
    if not recall:
        from vendor_intel.utils.domains import normalize_name

        def _vdom(v: dict[str, Any]) -> str:
            d = str(v.get("domain") or v.get("website") or "").strip().lower()
            return d.removeprefix("https://").removeprefix("http://").removeprefix("www.").split("/")[0]

        confirmed: list[dict[str, Any]] = []
        for v in final:
            ok, reason = _verification_status(v)
            if ok:
                confirmed.append(v)
            else:
                unverified_companies.append(
                    {
                        "company": v.get("company") or v.get("brand") or "",
                        "domain": _vdom(v),
                        "role": v.get("role") or "",
                        "reason": reason,
                    }
                )
        final = confirmed

        seed_name_set = {normalize_name(n).lower() for n in user_seed_names}
        final_doms = {_vdom(v) for v in final}
        listed = {(u["company"].lower(), u["domain"]) for u in unverified_companies}
        for r in export_rejected:
            nm = str(r.get("company") or "")
            if normalize_name(nm).lower() in seed_name_set and _vdom(r) not in final_doms:
                key = (nm.lower(), _vdom(r))
                if key not in listed:
                    unverified_companies.append(
                        {
                            "company": nm,
                            "domain": _vdom(r),
                            "role": r.get("role") or "",
                            "reason": _friendly_reject(
                                str(r.get("export_reject") or r.get("reject_reason") or "")
                            ),
                        }
                    )
                    listed.add(key)
        for nm in unresolved_seed_names:
            unverified_companies.append(
                {
                    "company": nm,
                    "domain": "",
                    "role": "",
                    "reason": "could not find an official website for this company",
                }
            )
        if unverified_companies:
            print(
                f"  [pipeline] {len(unverified_companies)} named/known companies could NOT be "
                f"confirmed — listed separately with reasons",
                flush=True,
            )

    elapsed = time.perf_counter() - t0
    meter = get_meter()
    llm_summary = meter.summary(
        provider=str(getattr(settings, "llm_provider", "opencode")),
        model=str(getattr(settings, "opencode_model", "")),
    )

    print(
        f"  [pipeline] {len(final)} exported (quality) of {len(classified)} classified "
        f"— no row padding",
        flush=True,
    )
    if not recall and len(final) < max(8, export_min or 0):
        print(
            f"  [pipeline] Thin result set ({len(final)} rows): discovery found limited "
            f"vetted companies; not inflating with low-quality fillers",
            flush=True,
        )
    if not recall and export_rejected:
        print(f"  [pipeline] {len(export_rejected)} rows dropped by quality_export", flush=True)

    # GUARANTEE: every curated seed lands in the final dataset, in its assigned section
    # (regardless of crawl/classification). Done before the role pass so they get summaries/roles.
    if not recall and not seed_audit:
        restored = _guarantee_seeds(final, unverified_companies, query_context, seed_domain_map)
        if restored:
            print(f"  [pipeline] guaranteed {restored} curated seed(s) into the final dataset", flush=True)

    # Clear market-specific role + market-alignment functionality for each company (one light
    # LLM pass). Companies the pass flags as clearly off-market are dropped from the export.
    if not recall and getattr(settings, "pipeline_market_roles", True):
        try:
            from vendor_intel.clients.claude import ClaudeClient
            from vendor_intel.pipeline.role_labels import (
                assign_market_roles,
                detect_multi_segments,
                is_excluded_segment,
                is_offmarket,
                place_unverified_in_sections,
            )

            client = ClaudeClient(settings)
            # Classify the not-yet-verified companies into their correct sections and fold them
            # into the main dataset (no separate "Not Verified" section).
            placed = place_unverified_in_sections(
                unverified_companies, final, query_context, settings, client
            )
            if placed:
                final.extend(placed)
            if getattr(client, "available", False):
                unverified_companies = []
            labelled = assign_market_roles(final, query_context, settings, client)
            # protect only the analyst's CURATED seed-file must-haves (not LLM-enumerated
            # "seeds", which can legitimately be off-market for this query)
            _curated_keys = {_nkey(n) for n in (query_context.get("seed_sections") or {})}

            def _is_curated(c: dict) -> bool:
                k = _nkey(c.get("company") or c.get("brand"))
                return any(
                    k and (k == cu or (len(k) >= 5 and k in cu) or (len(cu) >= 5 and cu in k))
                    for cu in _curated_keys
                )

            # Brief sizing-scope: companies the role pass flagged as belonging to an EXCLUDED
            # segment. Curated seeds are NEVER dropped (every seed must stay in the data); only
            # non-seed companies are pruned, conservatively (cap a couple). Everything flagged is
            # reported in the terminal so nothing leaves silently.
            exclude_active = bool(query_context.get("exclude_segments"))
            EXCLUDE_DROP_CAP = 2
            flagged_seed: list[dict] = []   # curated seeds in an excluded segment — always kept
            flagged_other: list[dict] = []  # non-seed, eligible for the capped prune

            dropped = 0
            for c in final:
                if not c.get("is_relevant"):
                    continue
                if exclude_active and is_excluded_segment(c):
                    (flagged_seed if _is_curated(c) else flagged_other).append(c)
                    continue
                if _is_curated(c):
                    continue
                if is_offmarket(c):
                    c["is_relevant"] = False
                    dropped += 1

            all_flagged = flagged_seed + flagged_other
            # capture the exclude reason before clearing labels (for the report)
            _why = {
                id(c): (str(c.get("market_role_detail") or "").strip() or "excluded segment")
                for c in all_flagged
            }
            excl_dropped = flagged_other[:EXCLUDE_DROP_CAP]
            for c in excl_dropped:
                c["is_relevant"] = False
            # kept-despite-flag (every seed + non-seeds beyond the cap): clear the 'Excluded-Segment'
            # label so they read as normal entries.
            for c in flagged_seed + flagged_other[EXCLUDE_DROP_CAP:]:
                c["market_role"] = ""
                c["market_role_detail"] = ""
            if all_flagged:
                print(
                    f"  [brief] {len(all_flagged)} company(ies) flagged outside the sizing scope "
                    f"(dropped {len(excl_dropped)}; {len(flagged_seed)} seed(s) always kept):",
                    flush=True,
                )
                for c in all_flagged:
                    nm = c.get("company") or c.get("brand") or ""
                    st = (
                        "DROPPED" if c in excl_dropped
                        else "KEPT (seed)" if c in flagged_seed
                        else "KEPT (cap reached - review for sizing)"
                    )
                    print(f"     - {nm}: {_why.get(id(c), '')} -> {st}", flush=True)

            multi = detect_multi_segments(final, query_context, settings, client)
            print(
                f"  [pipeline] market roles: {labelled} labelled"
                + (f", {dropped} off-market dropped" if dropped else "")
                + (f", {multi} multi-segment players" if multi else "")
                + (f", {len(placed)} unverified placed into sections" if placed else ""),
                flush=True,
            )
        except Exception as exc:
            print(f"  [pipeline] market-role labelling skipped: {exc}", flush=True)

    # Re-add the carried store companies (confirmed in a prior run, not re-processed this run).
    if store_carry:
        from vendor_intel.pipeline.quality_export import _export_domain as _xd

        _have = {_xd(c) for c in final}
        _added = 0
        for s in store_carry:
            dom = str(s.get("domain") or "").lower().removeprefix("www.")
            if not dom or dom in _have:
                continue
            rec = store_by_dom.get(dom, {})  # rich fields (role/section/summary) from the store
            _have.add(dom)
            _added += 1
            final.append({
                "company": s.get("name") or rec.get("name") or "", "domain": dom, "website": dom,
                "is_relevant": True, "confidence": 0.7, "is_seed": False,
                "discovery_source": "store",
                "role": rec.get("role") or "", "value_chain_section": rec.get("value_chain_section") or "",
                "company_summary": rec.get("company_summary") or "",
                "company_function": rec.get("company_function") or "",
            })
        if _added:
            print(f"  [pipeline] cost-saver: added {_added} carried store compan(ies) to output "
                  f"(no re-processing)", flush=True)

    print(f"  [pipeline] Total time: {elapsed / 60:.1f} min ({elapsed:.0f}s)", flush=True)
    print(
        f"  [pipeline] LLM calls: {llm_summary['llm_calls_total']} "
        f"(phase1={llm_summary['phase1_compile_calls']}, classify={llm_summary['classify_calls']}) "
        f"est. ${llm_summary['estimated_cost_usd']}",
        flush=True,
    )

    # Persist this run's confirmed (relevant) companies into the cross-run discovery store, so the
    # next run re-validates them and the cumulative union keeps climbing.
    if not recall:
        try:
            from vendor_intel.pipeline.discovery_store import enabled as _store_on
            from vendor_intel.pipeline.discovery_store import save_discovered

            if _store_on():
                _confirmed = [
                    {
                        "name": c.get("company") or c.get("brand") or "",
                        "domain": c.get("domain") or c.get("website") or "",
                        "role": c.get("role") or "",
                        "value_chain_section": c.get("value_chain_section") or "",
                        "company_summary": c.get("company_summary") or c.get("summary") or "",
                        "company_function": c.get("company_function") or "",
                    }
                    for c in final if c.get("is_relevant")
                ]
                n_store = save_discovered(store_market, _confirmed, country)
                print(f"  [pipeline] discovery store: {len(_confirmed)} confirmed this run -> "
                      f"{n_store} total in cumulative store", flush=True)
        except Exception as exc:
            print(f"  [pipeline] discovery store save skipped: {exc}", flush=True)

    recall_audit_md = ""
    _audit_stats: dict = {}
    if seed_audit and audit_seed_names:
        recall_audit_md = _print_seed_audit(
            audit_seed_names, final, classified, export_rejected, unverified_companies,
            market=str(scope.get("market") or query_context.get("industry") or ""),
            stats_out=_audit_stats,
        )

    return {
        "query": query,
        "recall_audit_md": recall_audit_md,
        "query_context": query_context,
        "pipeline_profile": profile,
        "recall_mode": recall,
        "use_ssc": use_ssc,
        "scope": scope,
        "completed_at": datetime.now(timezone.utc).isoformat(),
        "elapsed_seconds": round(elapsed, 1),
        "elapsed_minutes": round(elapsed / 60, 2),
        "llm_usage": llm_summary,
        "phase1_plan_path": plan_path,
        "phase2_path": p2.get("output_path"),
        "phase2_company_count": len(p2.get("companies") or []),
        "after_entity_gate": len(cleaned),
        "enriched_count": len(enriched),
        "classified_count": len(classified),
        "export_rejected_count": len(export_rejected) if not recall else 0,
        "relevant_companies": final,
        "seeds_pinned_total": seeds_pinned_total,
        "seeds_found_independently": seeds_found_independently,
        "seeds_found_independently_count": len(seeds_found_independently),
        "seed_audit_found": _audit_stats.get("found"),
        "seed_audit_total": _audit_stats.get("total"),
        "all_classified": classified,
        "export_rejected": export_rejected if not recall else [],
        "unverified_companies": unverified_companies,
    }


def _alpha_name(r: dict) -> str:
    from vendor_intel.utils.domains import fix_mojibake

    return fix_mojibake(str(r.get("company") or r.get("brand") or "")).strip().lower()


def _nkey(s: str) -> str:
    return re.sub(r"[^a-z0-9]", "", str(s or "").lower())


def _guarantee_seeds(
    final: list[dict], unverified: list[dict], query_context: dict, seed_domain_map: dict
) -> int:
    """Guarantee every curated seed appears in the final dataset, in its assigned section -
    regardless of crawl/classification outcome. Seeds in 'unverified' are moved into their
    section; seeds missing entirely are added as a row. Returns count added/restored."""
    seed_sections = query_context.get("seed_sections") or {}  # name -> assigned section
    if not seed_sections:
        return 0

    def _fuzzy(a: str, b: str) -> bool:
        return bool(a) and bool(b) and (a == b or (len(a) >= 5 and a in b) or (len(b) >= 5 and b in a))

    restored = 0
    for name, section in seed_sections.items():
        k = _nkey(name)
        if not k:
            continue
        existing = next((c for c in final if _fuzzy(k, _nkey(c.get("company") or c.get("brand")))), None)
        if existing is not None:
            # already in final — make sure it's kept and pinned to its section
            existing["is_relevant"] = True
            existing["is_seed"] = True
            if not str(existing.get("_forced_section") or "").strip():
                existing["_forced_section"] = section
            continue
        # not in final — add it (use a resolved domain if we have one)
        dom = seed_domain_map.get(k, "")
        final.append(
            {
                "company": name,
                "brand": name,
                "domain": dom,
                "website": dom,
                "role": "Manufacturer",
                "is_relevant": True,
                "is_seed": True,
                "is_major": True,
                "_forced_section": section,
                "confidence": 0.7,
                "quality_score": 0.7,
                "data_sources": "curated seed (guaranteed)",
            }
        )
        restored += 1
    # drop any guaranteed seeds out of the Not-Verified list (no double-listing)
    seed_keys = {_nkey(n) for n in seed_sections}
    unverified[:] = [
        u for u in unverified if not any(_fuzzy(_nkey(u.get("company")), sk) for sk in seed_keys)
    ]
    return restored


def _section_family(section_name: str) -> str | None:
    """Map a section heading to its value-chain family, or None if ambiguous."""
    s = (section_name or "").lower()
    # service / integration / connectivity sections (e.g. 'SATCOM Service Providers / Integrators')
    if any(w in s for w in ("service provider", "integrator", "integration", "connectivity", "managed service")):
        return "service"
    # network / capacity / operator sections
    if any(w in s for w in ("network", "capacity", "operator", "carrier")):
        return "operator"
    if any(w in s for w in ("distributor", "distribution", "trader", "wholesale")):
        return "distributor"
    if any(w in s for w in ("supplier", "raw material", "feedstock", "grower")):
        return "supplier"
    if any(w in s for w in ("manufacturer", "producer", "processing", "processor", "oem", "equipment",
                            "re-refin", "refining", "refiner")):
        return "manufacturer"
    if "brand" in s:
        return "brand"
    return None


_FAMILY_WORDS = {
    "distributor": ("distributor", "wholesal", "trader", "reseller", "retail", "marketer", "supplier", "brand", "offtaker"),
    "supplier": ("supplier", "grower", "feedstock", "raw", "farm", "producer", "cultivat"),
    "manufacturer": ("manufacturer", "producer", "processor", "maker", "refiner", "oem", "extractor", "miller", "production", "equipment"),
    "brand": ("brand", "marketer", "retail"),
    "service": ("service", "integrator", "integration", "provider", "managed", "solution", "aggregator", "reseller", "connectivity", "operator"),
    "operator": ("operator", "network", "capacity", "carrier", "fleet", "constellation", "provider", "infrastructure"),
}
_FAMILY_LABEL = {
    "distributor": "Distributor",
    "supplier": "Supplier",
    "manufacturer": "Producer",
    "brand": "Brand",
    "service": "Service Provider",
    "operator": "Operator",
}


def _consistent_role(market_role: str, section_name: str, main_product: str) -> str:
    """Keep the LLM role if it fits the section's family; otherwise replace it with a
    section-appropriate label (so a Distributors section never shows 'Processor')."""
    role = str(market_role or "").strip()
    fam = _section_family(section_name)
    if not fam:
        return role  # custom/ambiguous section — trust the LLM role
    if role and any(w in role.lower() for w in _FAMILY_WORDS[fam]):
        return role  # already consistent
    prod = (main_product or "").strip()
    label = _FAMILY_LABEL[fam]
    return f"{prod} {label}".strip() if prod else label


def _print_seed_audit(
    seed_names: list[str],
    final: list[dict],
    classified: list[dict],
    export_rejected: list[dict],
    unverified: list[dict],
    market: str = "",
    stats_out: dict | None = None,
) -> str:
    """Independent-recall audit: with the research-team list NOT pinned, report which of their
    companies the system re-discovered on its own (found / kept / rejected / missed, with reasons)
    and how many NET-NEW companies it surfaced beyond their list. Prints to console AND returns a
    CEO-facing markdown report string."""
    import re

    def _k(s: str) -> str:
        return re.sub(r"[^a-z0-9]", "", str(s or "").lower())

    # Connector + legal-form words that must NOT block a name match
    # ("Communications & Power Industries (CPI)" == "Communications and Power Industries CPI").
    _NOISE_TOKENS = frozenset(
        {
            "the", "and", "of", "for", "a",
            "inc", "incorporated", "ltd", "limited", "llc", "llp", "corp", "corporation",
            "co", "company", "group", "holding", "holdings", "plc", "sa", "sas", "ag",
            "gmbh", "spa", "pte", "pvt", "srl", "bv", "nv", "ab", "as", "oy", "kk",
        }
    )
    # Tokens too generic to anchor a match on their own (avoid false positives).
    _GENERIC_TOKENS = frozenset(
        {"satellite", "satellites", "satcom", "communications", "communication",
         "space", "systems", "system", "technologies", "technology", "networks",
         "network", "solutions", "global", "international", "telecom", "telecommunications"}
    )

    def _toks(s: str) -> set[str]:
        s = str(s or "").lower().replace("&", " and ")
        s = re.sub(r"\(.*?\)", " ", s)  # drop parentheticals like "(CPI)"
        return {t for t in re.split(r"[^a-z0-9]+", s) if t and t not in _NOISE_TOKENS}

    def _match(name: str, pool: list[tuple[str, dict]]) -> dict | None:
        k = _k(name)
        nt = _toks(name)
        if not k or not nt:
            return None
        nt_distinct = nt - _GENERIC_TOKENS
        for disp, rec in pool:
            d = _k(disp)
            if not d:
                continue
            # 1) joined-alnum substring (handles "STEngineering" vs "STEngineeringiDirect")
            if d == k or (len(k) >= 5 and k in d) or (len(d) >= 5 and d in k):
                return rec
            # 2) token-set match, ignoring connector/legal words & parentheticals
            dt = _toks(disp)
            if not dt:
                continue
            if nt == dt:
                return rec
            inter = nt & dt
            # subset match only when the shared core carries a DISTINCTIVE (non-generic) token,
            # so "SES" never swallows "SES Water" but "CPI" core still matches.
            if (nt <= dt or dt <= nt) and (inter & nt_distinct):
                return rec
        return None

    def _pool(rows: list[dict], name_key: str = "company") -> list[tuple[str, dict]]:
        return [((c.get(name_key) or c.get("brand") or ""), c) for c in rows]

    final_relevant = [c for c in final if c.get("is_relevant")]
    final_pool = _pool(final_relevant)
    rej_pool = _pool(export_rejected)
    unv_pool = _pool(unverified)
    allc_pool = _pool(classified)

    kept, rejected, unverif, found_not_kept, missed = [], [], [], [], []
    matched_records: list[dict] = []
    for s in seed_names:
        if (rec := _match(s, final_pool)) is not None:
            kept.append(s)
            matched_records.append(rec)
        elif (r := _match(s, rej_pool)) is not None:
            rejected.append((s, str(r.get("reject_reason") or r.get("reason") or "dropped by quality gate")))
        elif (u := _match(s, unv_pool)) is not None:
            unverif.append((s, str(u.get("reason") or "could not confirm")))
        elif _match(s, allc_pool):
            found_not_kept.append(s)
        else:
            missed.append(s)

    # NET-NEW: companies the system exported that are NOT on the research-team benchmark list.
    matched_ids = {id(r) for r in matched_records}
    net_new = [
        (c.get("company") or c.get("brand") or "") for c in final_relevant if id(c) not in matched_ids
    ]
    net_new = [n for n in net_new if n]

    n = len(seed_names)
    found = n - len(missed)
    pct = round(100 * found / n) if n else 0
    kept_pct = round(100 * len(kept) / n) if n else 0
    if stats_out is not None:
        stats_out.update({
            "found": found, "total": n, "kept": len(kept),
            "rejected": len(rejected), "missed": len(missed), "net_new": len(net_new),
        })

    print("\n" + "=" * 64, flush=True)
    print("  INDEPENDENT RECALL AUDIT (benchmark list NOT pinned)", flush=True)
    print("=" * 64, flush=True)
    print(f"  research-team companies     : {n}", flush=True)
    print(f"  re-discovered on its own    : {found}/{n} ({pct}%)", flush=True)
    print(f"    -> KEPT in final output   : {len(kept)}", flush=True)
    print(f"    -> REJECTED (with reason) : {len(rejected)}", flush=True)
    print(f"    -> UNVERIFIED (with reason): {len(unverif)}", flush=True)
    print(f"    -> found, not exported    : {len(found_not_kept)}", flush=True)
    print(f"  MISSED (not found at all)   : {len(missed)}", flush=True)
    print(f"  NET-NEW beyond their list   : {len(net_new)}", flush=True)
    # Provenance: HOW the system re-discovered each matched seed (web/structured vs LLM recall).
    _src_label = {
        "wikidata": "Wikidata (structured/web)",
        "search": "web search",
        "phase1_seed": "LLM enumeration / directory",
    }
    src_counts: dict[str, int] = {}
    for rec in matched_records:
        s = str(rec.get("discovery_source") or "search")
        src_counts[s] = src_counts.get(s, 0) + 1
    if src_counts:
        print("  -- re-discovered BY SOURCE (web/structured vs LLM):", flush=True)
        for s, ct in sorted(src_counts.items(), key=lambda x: -x[1]):
            print(f"       {_src_label.get(s, s):32s}: {ct}", flush=True)
    if missed:
        print("  -- MISSED:", flush=True)
        for s in missed:
            print(f"       x {s}", flush=True)
    print("=" * 64 + "\n", flush=True)

    # ---- CEO-facing markdown report ----
    def _bullets(items: list[str]) -> str:
        return "\n".join(f"- {x}" for x in items) if items else "_(none)_"

    md: list[str] = []
    title = market.strip() or "Market"
    md.append(f"# Independent Recall Audit — {title}\n")
    md.append(
        f"The research team identified **{n} companies** in this market. To test the system's *own* "
        f"discovery power, **their list was withheld from the run** — the engine received only the "
        f"market definition and had to find these companies through its AI market knowledge and "
        f"multi-source web search.\n"
    )
    md.append(
        f"## Result: the system independently re-discovered **{found} of {n} ({pct}%)**\n"
    )
    md.append("| Outcome | Companies |")
    md.append("|---|---:|")
    md.append(f"| Re-discovered and exported | {len(kept)} ({kept_pct}%) |")
    md.append(f"| Found, held for manual review (unverified) | {len(unverif)} |")
    md.append(f"| Found, filtered by the quality gate | {len(rejected)} |")
    md.append(f"| Found, not exported | {len(found_not_kept)} |")
    md.append(f"| Not found by the engine | {len(missed)} |")
    md.append("")
    md.append(
        f"Beyond the benchmark, the system also surfaced **{len(net_new)} additional companies** that "
        f"were not on the research team's list.\n"
    )
    md.append("### Re-discovered and exported\n" + _bullets(kept) + "\n")
    if unverif:
        md.append("### Found, held for review\n" + _bullets([f"{s} — {why}" for s, why in unverif]) + "\n")
    if rejected:
        md.append("### Found, filtered by quality gate\n" + _bullets([f"{s} — {why}" for s, why in rejected]) + "\n")
    if found_not_kept:
        md.append("### Found, not exported\n" + _bullets(found_not_kept) + "\n")
    md.append("### Not found by the engine\n" + _bullets(missed) + "\n")
    md.append("### Net-new companies (beyond the research team's list)\n" + _bullets(net_new) + "\n")
    return "\n".join(md)


def _alpha_sort_sections(grouped: list) -> list:
    """Order companies strictly A->Z within each section (one clean alphabetical list)."""
    return [(name, sorted(rows, key=_alpha_name)) for name, rows in grouped]


def _multi_segs(r: dict) -> list:
    return [s for s in (r.get("multi_segments") or []) if isinstance(s, dict) and str(s.get("role") or "").strip()]


def _partition_multisegment(rows: list) -> tuple[list, list]:
    """Split into (multi-segment players, single). Multi = operates in 2+ sections; placed in a
    single block (appended at the bottom by the exporters) ordered most-segments-first then A->Z.
    Multi companies appear ONCE."""
    multi = [r for r in rows if len(_multi_segs(r)) >= 2]
    multi_ids = {id(r) for r in multi}
    single = [r for r in rows if id(r) not in multi_ids]
    multi.sort(key=lambda r: (-len(_multi_segs(r)), _alpha_name(r)))
    return multi, single


def _multi_role_label(r: dict) -> str:
    """Join a multi-segment company's distinct roles with ' + ' (e.g. 'Re-Refiner + Offtaker')."""
    return " + ".join(dict.fromkeys(str(s["role"]).strip() for s in _multi_segs(r)))


def save_pipeline_csv(result: dict[str, Any], path: str) -> str:
    """CSV with sources, website, quality score."""
    import csv
    from pathlib import Path

    ctx = result.get("query_context") or {}
    industry = str(ctx.get("industry") or result.get("query") or "")
    scope = ctx.get("scope") if isinstance(ctx.get("scope"), dict) else result.get("scope")
    country = str(ctx.get("country") or "global")
    rows = [
        r for r in (result.get("relevant_companies") or []) if r.get("is_relevant")
    ]
    from vendor_intel.pipeline.quality_export import dedupe_export_rows

    rows = dedupe_export_rows(rows)

    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)

    from vendor_intel.pipeline.sections import (
        build_section_taxonomy,
        group_into_sections,
        main_product_label,
    )

    _scope = scope if isinstance(scope, dict) else None
    main_product = main_product_label(ctx, _scope)
    custom_sections = [
        str(s).strip() for s in (ctx.get("sections") or []) if str(s).strip()
    ]
    multi_rows, single_rows = _partition_multisegment(rows)
    if custom_sections:
        taxonomy = custom_sections
        grouped = group_into_sections(single_rows, taxonomy, main_product, custom=True)
    else:
        taxonomy = build_section_taxonomy(main_product)
        grouped = group_into_sections(single_rows, taxonomy, main_product)
    grouped = _alpha_sort_sections(grouped)
    if multi_rows:
        grouped = grouped + [("Multi-Segment Players", multi_rows)]

    _HEADER = [
        "#",
        "Company",
        "Brand",
        "Parent_or_Independent",
        "Website",
        "Functionality",
        "Geography",
        "Is_Relevant",
        "Summary",
    ]

    def _row_cells(i: int, r: dict, section_name: str = "") -> list:
        from vendor_intel.pipeline.csv_fields import (
            company_summary_cell,
            format_csv_functionality,
        )

        from vendor_intel.utils.domains import fix_mojibake

        kp_raw = r.get("key_products", "") or r.get("product_system_types", "")
        role = str(r.get("role") or "Other")
        role_desc = str(r.get("role_description", "") or r.get("role_detail", ""))
        company = fix_mojibake(str(r.get("company", "")))
        brand = fix_mojibake(str(r.get("brand", "") or "")) or company
        # Functionality is composed from the market-role detail when the role pass ran,
        # else from the generic classifier role + key products.
        mr_detail = str(r.get("market_role_detail") or "").strip()
        functionality = mr_detail or format_csv_functionality(
            role,
            role_desc,
            kp_raw,
            market=industry,
            company_function=str(r.get("company_function") or ""),
            scope=_scope,
        )
        return [
            i,
            company,
            brand,
            r.get("parent", "") or r.get("parent_or_independent", ""),
            r.get("website", "") or r.get("domain", ""),
            functionality,
            country,
            "yes" if r.get("is_relevant") else "no",
            company_summary_cell(
                str(r.get("company_summary") or ""),
                company=str(r.get("company") or ""),
                role=role,
                role_description=role_desc,
                key_products=kp_raw,
                max_sentences=4,
                max_len=460,
            ),
        ]

    unverified = sorted(
        result.get("unverified_companies") or [],
        key=lambda u: str(u.get("company") or "").lower(),
    )

    def _write_csv(target: Path) -> None:
        from vendor_intel.utils.domains import fix_mojibake

        with target.open("w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            w.writerow(_HEADER)
            for sect_idx, (section_name, section_rows) in enumerate(grouped):
                if sect_idx > 0:
                    w.writerow([])  # blank spacer row between sections
                w.writerow([f"=== {section_name} ({len(section_rows)}) ==="])
                for i, r in enumerate(section_rows, 1):
                    w.writerow(_row_cells(i, r, section_name))
            if unverified:
                w.writerow([])
                w.writerow([f"=== Not Verified - Could Not Confirm ({len(unverified)}) ==="])
                for i, u in enumerate(unverified, 1):
                    w.writerow(
                        [
                            i,
                            fix_mojibake(str(u.get("company") or "")),
                            "", "",
                            u.get("domain", ""),
                            "(unverified)",
                            country, "no",
                            str(u.get("reason") or ""),
                        ]
                    )

    try:
        _write_csv(out)
    except PermissionError:
        alt = out.with_name(
            f"{out.stem}_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}{out.suffix}"
        )
        _write_csv(alt)
        print(
            f"  [pipeline] CSV locked ({out.name}) — saved: {alt.resolve()}",
            flush=True,
        )
        return str(alt.resolve())

    print(f"  [pipeline] CSV saved: {out.resolve()}", flush=True)
    return str(out.resolve())


def save_pipeline_xlsx(result: dict[str, Any], path: str) -> str:
    """Polished Excel deliverable: title banner, frozen styled header, bold shaded section
    bands, numbered companies, clickable website links. No 'Industry' column."""
    from pathlib import Path

    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    from vendor_intel.pipeline.csv_fields import (
        company_summary_cell,
        format_csv_functionality,
    )
    from vendor_intel.pipeline.quality_export import dedupe_export_rows
    from vendor_intel.pipeline.sections import (
        build_section_taxonomy,
        group_into_sections,
        main_product_label,
    )
    from vendor_intel.utils.domains import fix_mojibake

    ctx = result.get("query_context") or {}
    industry = str(ctx.get("industry") or result.get("query") or "")
    scope = ctx.get("scope") if isinstance(ctx.get("scope"), dict) else result.get("scope")
    _scope = scope if isinstance(scope, dict) else None
    country = str(ctx.get("country") or "global")
    rows = dedupe_export_rows(
        [r for r in (result.get("relevant_companies") or []) if r.get("is_relevant")]
    )

    main_product = main_product_label(ctx, _scope)
    custom_sections = [str(s).strip() for s in (ctx.get("sections") or []) if str(s).strip()]
    if custom_sections:
        # Brief-scoped Excel: only the sections named in the brief (no auto-taxonomy, drop the
        # catch-all 'Other' of off-brief rows), with 'Multi-Segment Players' appended at the bottom.
        multi_rows, single_rows = _partition_multisegment(rows)
        grouped = group_into_sections(single_rows, custom_sections, main_product, custom=True)
        allowed = {s.strip().lower() for s in custom_sections}
        grouped = _alpha_sort_sections(
            [(n, r) for n, r in grouped if n.strip().lower() in allowed]
        )
        if multi_rows:
            grouped = grouped + [("Multi-Segment Players", multi_rows)]
    else:
        multi_rows, single_rows = _partition_multisegment(rows)
        grouped = _alpha_sort_sections(
            group_into_sections(single_rows, build_section_taxonomy(main_product), main_product)
        )
        if multi_rows:
            grouped = grouped + [("Multi-Segment Players", multi_rows)]

    header = ["#", "Company", "Brand", "Parent_or_Independent", "Website",
              "Functionality", "Geography", "Is_Relevant", "Summary"]
    ncol = len(header)
    last_col = get_column_letter(ncol)

    NAVY = PatternFill("solid", fgColor="1F3864")
    BAND = PatternFill("solid", fgColor="D6E0F0")
    white_bold = Font(bold=True, color="FFFFFF")
    band_font = Font(bold=True, color="1F3864", size=12)
    wrap = Alignment(vertical="top", wrap_text=True)

    wb = Workbook()
    ws = wb.active
    ws.title = (main_product or industry or "Companies")[:31]

    # Title banner + meta
    ws.merge_cells(f"A1:{last_col}1")
    t = ws["A1"]; t.value = industry or main_product or "Market Landscape"
    t.font = Font(bold=True, size=16, color="1F3864")
    ws.merge_cells(f"A2:{last_col}2")
    m = ws["A2"]; m.value = f"Geography: {country}     Companies: {len(rows)}"
    m.font = Font(italic=True, color="595959")

    # Header row (row 3)
    hdr = 3
    for c, name in enumerate(header, 1):
        cell = ws.cell(row=hdr, column=c, value=name)
        cell.fill = NAVY; cell.font = white_bold
        cell.alignment = Alignment(vertical="center")
    ws.freeze_panes = f"A{hdr + 1}"

    rownum = hdr + 1
    for section_name, section_rows in grouped:
        ws.merge_cells(f"A{rownum}:{last_col}{rownum}")
        b = ws.cell(row=rownum, column=1, value=f"{section_name}  ({len(section_rows)})")
        b.fill = BAND; b.font = band_font
        rownum += 1
        for i, r in enumerate(section_rows, 1):
            role = str(r.get("role") or "Other")
            kp_raw = r.get("key_products", "") or r.get("product_system_types", "")
            role_desc = str(r.get("role_description", "") or r.get("role_detail", ""))
            company = fix_mojibake(str(r.get("company", "")))
            brand = fix_mojibake(str(r.get("brand", "") or "")) or company
            func = str(r.get("market_role_detail") or "").strip() or format_csv_functionality(
                role, role_desc, kp_raw, market=industry,
                company_function=str(r.get("company_function") or ""), scope=_scope,
            )
            website = str(r.get("website", "") or r.get("domain", "") or "")
            summary = company_summary_cell(
                str(r.get("company_summary") or ""), company=company, role=role,
                role_description=role_desc, key_products=kp_raw, max_sentences=4, max_len=460,
            )
            vals = [i, company, brand,
                    r.get("parent", "") or r.get("parent_or_independent", ""),
                    website, func, country, "yes" if r.get("is_relevant") else "no", summary]
            for c, v in enumerate(vals, 1):
                cell = ws.cell(row=rownum, column=c, value=v)
                cell.alignment = wrap
            # clickable website
            if website:
                link = website if website.startswith("http") else f"https://{website}"
                wc = ws.cell(row=rownum, column=5)
                wc.hyperlink = link
                wc.font = Font(color="0563C1", underline="single")
            rownum += 1

    widths = [4, 26, 20, 22, 30, 40, 12, 10, 70]
    for c, wdt in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(c)].width = wdt

    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    try:
        wb.save(str(out))
    except PermissionError:
        out = out.with_name(
            f"{out.stem}_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}{out.suffix}"
        )
        wb.save(str(out))
    print(f"  [pipeline] XLSX saved: {out.resolve()}", flush=True)
    return str(out.resolve())


def save_pipeline_docx(result: dict[str, Any], path: str) -> str:
    """CEO Word doc: market heading → section sub-headings → numbered companies
    with each company's functionality in brackets. Mirrors the CSV sectioning."""
    from pathlib import Path

    from docx import Document
    from docx.shared import Pt

    from vendor_intel.pipeline.csv_fields import format_csv_functionality
    from vendor_intel.pipeline.quality_export import dedupe_export_rows
    from vendor_intel.pipeline.sections import (
        build_section_taxonomy,
        group_into_sections,
        main_product_label,
    )
    from vendor_intel.utils.domains import fix_mojibake

    ctx = result.get("query_context") or {}
    industry = str(ctx.get("industry") or result.get("query") or "")
    scope = ctx.get("scope") if isinstance(ctx.get("scope"), dict) else result.get("scope")
    _scope = scope if isinstance(scope, dict) else None
    country = str(ctx.get("country") or "global")
    rows = [r for r in (result.get("relevant_companies") or []) if r.get("is_relevant")]
    rows = dedupe_export_rows(rows)

    main_product = main_product_label(ctx, _scope)
    custom_sections = [str(s).strip() for s in (ctx.get("sections") or []) if str(s).strip()]
    multi_rows, single_rows = _partition_multisegment(rows)
    if custom_sections:
        grouped = group_into_sections(single_rows, custom_sections, main_product, custom=True)
    else:
        grouped = group_into_sections(single_rows, build_section_taxonomy(main_product), main_product)
    grouped = _alpha_sort_sections(grouped)
    if multi_rows:
        grouped = grouped + [("Multi-Segment Players", multi_rows)]

    doc = Document()
    doc.add_heading(industry or main_product or "Market Landscape", level=0)
    meta = doc.add_paragraph()
    meta.add_run(f"Geography: {country}     Companies: {len(rows)}").italic = True

    for section_name, section_rows in grouped:
        doc.add_heading(f"{section_name} ({len(section_rows)})", level=1)
        for i, r in enumerate(section_rows, 1):
            company = fix_mojibake(str(r.get("company") or r.get("brand") or "")).strip()
            role = str(r.get("role") or "Other")
            kp_raw = r.get("key_products", "") or r.get("product_system_types", "")
            # Functionality only in the brackets (no role label) — what the company does.
            func = str(r.get("market_role_detail") or "").strip() or format_csv_functionality(
                role,
                str(r.get("role_description", "") or r.get("role_detail", "")),
                kp_raw,
                market=industry,
                company_function=str(r.get("company_function") or ""),
                scope=_scope,
            )
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            p.add_run(f"{i}. ")
            p.add_run(company).bold = True
            if func:
                p.add_run(f" ({func})")

    unverified = sorted(
        result.get("unverified_companies") or [],
        key=lambda u: str(u.get("company") or "").lower(),
    )
    if unverified:
        doc.add_heading(f"Not Verified - Could Not Confirm ({len(unverified)})", level=1)
        note = doc.add_paragraph()
        note.add_run(
            "Named/known companies the system could not substantiate. Reason given per company."
        ).italic = True
        for i, u in enumerate(unverified, 1):
            company = fix_mojibake(str(u.get("company") or "")).strip()
            reason = str(u.get("reason") or "")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            p.add_run(f"{i}. ")
            p.add_run(company).bold = True
            if reason:
                p.add_run(f" - {reason}")

    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(out))
    except PermissionError:
        alt = out.with_name(
            f"{out.stem}_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}{out.suffix}"
        )
        doc.save(str(alt))
        print(f"  [pipeline] DOCX locked ({out.name}) — saved: {alt.resolve()}", flush=True)
        return str(alt.resolve())
    print(f"  [pipeline] DOCX saved: {out.resolve()}", flush=True)
    return str(out.resolve())


def run_pipeline_sync(
    query_context: dict[str, Any],
    settings: Settings | None = None,
    **kwargs: Any,
) -> dict[str, Any]:
    return asyncio.run(run_pipeline(query_context, settings, **kwargs))
